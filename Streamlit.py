import streamlit as st

# Set page config BEFORE anything else
st.set_page_config(
    page_title="Resume Parser", 
    page_icon="📄",
    layout="wide"
)

import logging
import subprocess
import sys
import re
from datetime import datetime
from typing import Dict, Optional
from io import BytesIO
from docx import Document
from pypdf import PdfReader
from docx.shared import Pt


# Install missing packages
REQUIRED_PACKAGES = {
    "python-docx": "docx",
    "pypdf": "pypdf"
}

missing = []
for pkg, imp in REQUIRED_PACKAGES.items():
    try:
        __import__(imp)
    except ImportError:
        missing.append(pkg)

if missing:
    for pkg in missing:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])
        except Exception as e:
            st.error(f"Please manually install {pkg} with pip install {pkg}")
            logging.error(f"Error installing {pkg}: {e}")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("parser.log"),
        logging.StreamHandler()
    ]
)

class ResumeParser:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.common_skills = {"python", "java", "javascript", "typescript", "html", "css", "sql", "nosql", 
                               "react", "angular", "vue", "node.js", "django", "flask", "express", "spring",
                               "docker", "kubernetes", "aws", "azure", "gcp", "git", "agile", "scrum", "jira",
                               "jenkins", "ci/cd", "rest api", "graphql", "mongodb", "mysql", "postgresql", 
                               "oracle", "data analysis", "machine learning", "deep learning", "ai", "nltk",
                               "pandas", "numpy", "tensorflow", "pytorch", "keras", "scikit-learn", "excel", 
                               "powerpoint", "word", "tableau", "power bi", "linux", "windows", "macos", 
                               "networking", "security"}

    def read_pdf(self, file) -> Optional[str]:
        try:
            reader = PdfReader(file)
            text = "".join([page.extract_text() or "" for page in reader.pages])
            return text
        except Exception as e:
            self.logger.error(f"PDF read error: {e}")
            return None

    def read_docx(self, file) -> Optional[str]:
        try:
            doc = Document(file)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            self.logger.error(f"DOCX read error: {e}")
            return None

    def read_text_file(self, file) -> Optional[str]:
        try:
            return file.read().decode('utf-8')
        except Exception as e:
            self.logger.error(f"Text read error: {e}")
            return None

    def extract_email(self, text: str) -> Optional[str]:
        match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        return match.group(0) if match else None

    def extract_phone(self, text: str) -> Optional[str]:
        match = re.search(r'(\+?\d{1,4}[-\s.]?)?(\(?\d{2,4}\)?[-\s.]?)?\d{3,5}[-\s.]?\d{4,6}', text)
        return match.group(0) if match else None

    def extract_skills(self, text: str) -> list:
        text = text.lower()
        return [skill for skill in self.common_skills if re.search(rf'\b{re.escape(skill)}\b', text)]

    def extract_education(self, text: str) -> list:
        patterns = [
            r'\b(Ph\\.?D\\.?|Doctor of Philosophy)\b',
            r'\b(M\\.?S\\.?|MBA|Master of [A-Za-z]+)\b',
            r'\b(B\\.?S\\.?|B\\.?A\\.?|Bachelor of [A-Za-z]+)\b',
            r'\b(Associate\'s? Degree|A\\.?A\\.?|A\\.?S\\.?)\b'
        ]
        matches = []
        for pat in patterns:
            for match in re.finditer(pat, text, re.IGNORECASE):
                start, end = max(0, match.start()-50), min(len(text), match.end()+50)
                matches.append(text[start:end].strip())
        return matches

    def create_word_report(self, data: Dict) -> BytesIO:
        try:
            import os
            import streamlit as st
            
            # Try multiple possible locations for the template
            possible_paths = [
                "template.docx",  # Current working directory
                os.path.join(os.path.dirname(os.path.abspath(__file__)), "template.docx"),  # Script directory
                os.path.abspath("template.docx"),  # Absolute path in current directory
                os.path.join(os.getcwd(), "template.docx")  # Explicit current working directory
            ]
            
            # Log all the paths we're trying
            template_found = False
            for path in possible_paths:
                self.logger.info(f"Checking for template at: {path}")
                if os.path.exists(path):
                    template_path = path
                    template_found = True
                    self.logger.info(f"Template found at: {template_path}")
                    
                    # Check file size to ensure it's not empty
                    file_size = os.path.getsize(template_path)
                    self.logger.info(f"Template file size: {file_size} bytes")
                    if file_size == 0:
                        self.logger.warning(f"Template file exists but is empty: {template_path}")
                        continue
                    
                    # Try to load the template
                    try:
                        self.logger.info(f"Attempting to load template from: {template_path}")
                        doc = Document(template_path)
                        self.logger.info("Successfully loaded template document")
                        break
                    except Exception as template_error:
                        self.logger.error(f"Error loading template from {template_path}: {template_error}")
                        continue
            else:
                # This block runs if no template was successfully loaded
                self.logger.warning("No template could be loaded, using basic document")
                template_found = False if template_found else False  # For clarity in logs
                doc = Document()
                
                # Create a basic report without template
                doc.add_heading("Resume Analysis Report", 0)
                
                doc.add_heading("Contact Information", level=1)
                doc.add_paragraph(f"Email: {data.get('email', 'N/A')}")
                doc.add_paragraph(f"Phone: {data.get('phone', 'N/A')}")
                
                doc.add_heading("Skills", level=1)
                skills_text = ", ".join(data.get("skills", [])) or "No common skills detected"
                doc.add_paragraph(skills_text)
                
                doc.add_heading("Education", level=1)
                for edu in data.get("education", []):
                    doc.add_paragraph(f"• {edu}", style="List Bullet")
                if not data.get("education"):
                    doc.add_paragraph("N/A")
                    
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                return buf
            
            # If we get here, we've successfully loaded the template
            self.logger.info("Processing template replacements")
            
            # Store original paragraphs for debugging
            original_paragraphs = [p.text for p in doc.paragraphs]
            self.logger.info(f"Template has {len(original_paragraphs)} paragraphs")
            self.logger.info(f"First few paragraphs: {original_paragraphs[:3]}")
            
            replacements = {
                "{{email}}": data.get("email", "N/A"),
                "{{phone}}": data.get("phone", "N/A"),
                "{{skills}}": ", ".join(data.get("skills", [])) or "No common skills detected",
                "{{education}}": "\n".join(f"• {edu}" for edu in data.get("education", [])) or "N/A"
            }
            
            # Log replacements for debugging
            self.logger.info(f"Replacements to be made: {replacements}")
            
            # Process all paragraphs to replace placeholders while preserving formatting
            for i, paragraph in enumerate(doc.paragraphs):
                # Check if paragraph contains any placeholder
                original_text = paragraph.text
                modified = False
                for key, val in replacements.items():
                    if key in original_text:
                        self.logger.info(f"Found placeholder {key} in paragraph {i}")
                        modified = True
                        
                        # Store formatting of runs
                        runs_formatting = []
                        for run in paragraph.runs:
                            runs_formatting.append({
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font_size': run.font.size if hasattr(run.font, 'size') and run.font.size else None,
                                'font_name': run.font.name if hasattr(run.font, 'name') and run.font.name else None,
                            })
                        
                        # Clear the paragraph
                        for j in range(len(paragraph.runs)-1, -1, -1):
                            p = paragraph._p
                            p.remove(paragraph.runs[j]._r)
                        
                        # Replace the placeholder in the text
                        new_text = original_text.replace(key, val)
                        
                        # Re-add the text while trying to preserve formatting
                        paragraph.add_run(new_text)
                        
                        # Try to restore formatting of the first run from original formatting
                        if runs_formatting:
                            for attr, value in runs_formatting[0].items():
                                if attr != 'text' and value is not None:
                                    if attr == 'font_size':
                                        paragraph.runs[0].font.size = value
                                    elif attr == 'font_name':
                                        paragraph.runs[0].font.name = value
                                    else:
                                        setattr(paragraph.runs[0], attr, value)
                
                if modified:
                    self.logger.info(f"Modified paragraph {i}: '{original_text}' -> '{paragraph.text}'")
            
            # Process any tables in the document
            table_count = len(doc.tables)
            self.logger.info(f"Template has {table_count} tables")
            
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p_idx, paragraph in enumerate(cell.paragraphs):
                            original_text = paragraph.text
                            modified = False
                            for key, val in replacements.items():
                                if key in original_text:
                                    modified = True
                                    paragraph.text = original_text.replace(key, val)
                            if modified:
                                self.logger.info(f"Modified text in table {t_idx}, row {r_idx}, cell {c_idx}: '{original_text}' -> '{paragraph.text}'")
            
            self.logger.info("Finished processing template, saving document")
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            self.logger.info("Document saved successfully")
            return buf
    
        except Exception as e:
            self.logger.error(f"Report creation failed: {e}")
            import traceback
            self.logger.error(f"Traceback: {traceback.format_exc()}")
            
            # Create a simple emergency report
            doc = Document()
            doc.add_heading("Resume Analysis - Error Recovery", 0)
            doc.add_paragraph(f"Error creating formatted report: {str(e)}")
            doc.add_paragraph(f"Email: {data.get('email', 'N/A')}")
            doc.add_paragraph(f"Phone: {data.get('phone', 'N/A')}")
            
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf



@st.cache_resource
def get_parser():
    return ResumeParser()

def main():
    #st.set_page_config(page_title="Resume Parser", page_icon="📄", layout="wide")
    st.title("📄 Resume Parser App")
    st.write("Upload a resume (PDF, DOCX, TXT) to extract contact info, skills, and education.")

    uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx", "txt"])


    # Add this somewhere in your main() function
    if st.checkbox("Debug mode"):
        import os
        st.subheader("Path Debugging Information")
        st.write(f"Current working directory: {os.getcwd()}")
        st.write(f"__file__ value: {__file__}")
        st.write(f"Script directory: {os.path.dirname(os.path.abspath(__file__))}")
        
        # List files in current directory
        st.write("Files in current directory:")
        files = os.listdir(os.getcwd())
        st.write(files)
        
        # Check if template exists in various locations
        locations = [
            "template.docx",
            os.path.join(os.path.dirname(os.path.abspath(__file__)), "template.docx"),
            os.path.abspath("template.docx"),
            os.path.join(os.getcwd(), "template.docx")
        ]
        
        for loc in locations:
            st.write(f"Template at {loc}: {'Exists' if os.path.exists(loc) else 'Not found'}")


    
    if uploaded_file:
        st.info(f"File uploaded: {uploaded_file.name}")
        parser = get_parser()
        ext = uploaded_file.name.split('.')[-1].lower()

        if ext == "pdf":
            text = parser.read_pdf(uploaded_file)
        elif ext == "docx":
            text = parser.read_docx(uploaded_file)
        elif ext == "txt":
            text = parser.read_text_file(uploaded_file)
        else:
            st.error("Unsupported file format")
            return

        if not text:
            st.error("Could not read the file.")
            return

        data = {
            "email": parser.extract_email(text),
            "phone": parser.extract_phone(text),
            "skills": parser.extract_skills(text),
            "education": parser.extract_education(text),
            "raw_text": text
        }

        st.success("Resume parsed successfully!")

        col1, col2 = st.columns(2)
        with col1:
            st.subheader("📇 Contact Info")
            st.write(f"📧 {data['email'] or 'Not found'}")
            st.write(f"📱 {data['phone'] or 'Not found'}")
        with col2:
            st.subheader("🧠 Skills")
            if data['skills']:
                for skill in data['skills']:
                    st.markdown(f"- {skill}")
            else:
                st.write("No common skills found")

        if data['education']:
            st.subheader("🎓 Education")
            for edu in data['education']:
                st.write(f"• {edu}")

        st.subheader("📥 Download Report")
        doc = parser.create_word_report(data)
        st.download_button(
            label="Download Word Report",
            data=doc,
            file_name=f"resume_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        with st.expander("View Raw Resume Text"):
            st.code(text, language='text')

if __name__ == "__main__":
    main()
